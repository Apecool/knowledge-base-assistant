"""
File Parser Service — Parses uploaded documents (.txt, .md, .pdf, .docx)
Extracts title and content text for knowledge base indexing.
"""
import os
import re
from typing import Dict, Optional, Tuple
from pathlib import Path


class FileParser:
    """
    Parses uploaded files and extracts readable content.
    Supports: .txt, .md, .pdf, .docx
    """

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}

    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if a file extension is supported."""
        ext = Path(filename).suffix.lower()
        return ext in FileParser.SUPPORTED_EXTENSIONS

    @staticmethod
    def extract_title(filename: str, content: str) -> str:
        """
        Extract a title from the file content or fall back to filename.
        Priority:
        1. First # heading in markdown
        2. First non-empty line
        3. Filename without extension
        """
        # Check for markdown heading
        heading_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        if heading_match:
            return heading_match.group(1).strip()

        # Check first non-empty line
        for line in content.split('\n'):
            line = line.strip()
            if line:
                return line[:120]  # Cap at 120 chars

        # Fallback to filename
        return Path(filename).stem

    @staticmethod
    def parse_txt(content_bytes: bytes) -> Tuple[str, str]:
        """Parse a plain text file."""
        content = content_bytes.decode('utf-8', errors='replace')
        title = FileParser.extract_title("document.txt", content)
        return title, content

    @staticmethod
    def parse_md(content_bytes: bytes) -> Tuple[str, str]:
        """Parse a markdown file - keep formatting as-is."""
        content = content_bytes.decode('utf-8', errors='replace')
        title = FileParser.extract_title("document.md", content)
        return title, content

    @staticmethod
    def parse_pdf(content_bytes: bytes) -> Tuple[str, str]:
        """Parse a PDF file using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content_bytes, filetype="pdf")
            content_parts = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    content_parts.append(text.strip())
            doc.close()
            content = "\n\n".join(content_parts)
            title = FileParser.extract_title("document.pdf", content)
            return title, content
        except ImportError:
            raise RuntimeError("pymupdf not installed. Run: pip install pymupdf")
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def parse_docx(content_bytes: bytes) -> Tuple[str, str]:
        """Parse a Word document using python-docx."""
        try:
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(content_bytes))
            content_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text.strip())

            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    content_parts.append(" | ".join(cells))

            content = "\n\n".join(content_parts)
            title = FileParser.extract_title("document.docx", content)
            return title, content
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {str(e)}")

    @classmethod
    def parse(cls, filename: str, content_bytes: bytes) -> Dict:
        """
        Parse an uploaded file and return structured content.

        Args:
            filename: Original filename (used to detect type and derive title)
            content_bytes: Raw file bytes

        Returns:
            Dict with keys: title, content, file_type, file_size
        """
        ext = Path(filename).suffix.lower()

        if ext == ".txt":
            title, content = cls.parse_txt(content_bytes)
        elif ext == ".md":
            title, content = cls.parse_md(content_bytes)
        elif ext == ".pdf":
            title, content = cls.parse_pdf(content_bytes)
        elif ext == ".docx":
            title, content = cls.parse_docx(content_bytes)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {cls.SUPPORTED_EXTENSIONS}")

        return {
            "title": title,
            "content": content,
            "file_type": ext.lstrip('.'),
            "file_size": len(content_bytes),
        }